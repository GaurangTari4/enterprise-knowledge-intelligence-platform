from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("output/docx/employee_handbook_sample.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        element = tc_mar.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9DEE7")


def set_table_fixed_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_table_widths(table, widths_in_inches: list[float]) -> None:
    table.autofit = False
    set_table_fixed_layout(table)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in_inches):
            cell.width = Inches(width)
            set_cell_margins(cell)
    set_table_borders(table)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_paragraph_with_run(doc, text: str, *, bold=False, italic=False, size=None, color=None, font="Calibri"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = font
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    return p


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after, color in [
        ("Heading 1", 16, 18, 10, "2E74B5"),
        ("Heading 2", 13, 14, 7, "2E74B5"),
        ("Heading 3", 12, 10, 5, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Employee Handbook")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F3A5F")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle_run = subtitle.add_run("Sample template for internal review - based on public EEOC, DOL, and OSHA guidance")
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor.from_string("5A6573")
    subtitle_run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(10)
    note_run = note.add_run(
        "Important: This handbook is a sample drafting aid, not legal advice. Replace placeholders, align it with company policy, and have local counsel review it before adoption."
    )
    note_run.font.name = "Calibri"
    note_run.font.size = Pt(10.5)
    note_run.font.color.rgb = RGBColor.from_string("1F3A5F")
    note_run.bold = True
    note_run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    note_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def add_at_a_glance_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [1.8, 4.7])
    hdr = table.rows[0].cells
    hdr[0].text = "Topic"
    hdr[1].text = "Sample instruction"
    set_repeat_table_header(table.rows[0])
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True

    rows = [
        ("Employment relationship", "Employment is at-will unless a signed agreement says otherwise."),
        ("Equal opportunity", "Hiring, promotion, discipline, and pay decisions are based on job-related factors."),
        ("Harassment and retaliation", "Report concerns early; the company will investigate promptly and no one may retaliate."),
        ("Pay and timekeeping", "Record all hours worked accurately and report overtime before it is worked when possible."),
        ("Annual leave", "Full-time employees accrue 20 days of annual leave per year."),
        ("Safety and security", "Stop unsafe work, report hazards immediately, and protect company data and equipment."),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)


def add_section_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    return p


def add_bullets(doc: Document, bullets: list[str]) -> None:
    for bullet in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(bullet)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def add_source_entry(doc: Document, title: str, url: str, note: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{title}: ")
    r1.bold = True
    r2 = p.add_run(f"{note} {url}")
    for run in (r1, r2):
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def build_document() -> Document:
    doc = Document()
    style_document(doc)
    add_title_block(doc)
    add_at_a_glance_table(doc)

    add_section_heading(doc, "1. Purpose and scope", 1)
    doc.add_paragraph(
        "This handbook gives employees a clear, practical summary of workplace expectations. It is designed as a sample template for a private employer in the United States and should be customized for the company, state, city, and industry before it is issued."
    )
    doc.add_paragraph(
        "Use this handbook as a reference for day-to-day decisions, reporting concerns, and understanding basic rights and responsibilities. When this handbook and a written offer letter or signed agreement conflict, the signed agreement controls to the extent allowed by law."
    )

    add_section_heading(doc, "2. Employment relationship", 1)
    add_bullets(
        doc,
        [
            "Employment is at-will unless a signed agreement or local law says otherwise.",
            "Nothing in this handbook creates a contract of employment or guarantees a specific length of employment.",
            "The company may change policies, schedules, benefits, or procedures when business needs or the law require it.",
        ],
    )

    add_section_heading(doc, "3. Equal opportunity, harassment, and retaliation", 1)
    doc.add_paragraph(
        "The company is committed to fair treatment and does not tolerate discrimination, harassment, or retaliation. Employment decisions are based on job-related qualifications, business needs, and lawful requirements."
    )
    add_bullets(
        doc,
        [
            "Protected characteristics include race, color, religion, sex, pregnancy, sexual orientation, gender identity, national origin, age, disability, and genetic information where applicable.",
            "Harassment includes unwelcome conduct that is severe or pervasive enough to create a hostile work environment or that affects a term or condition of employment.",
            "Retaliation against anyone who raises a concern or participates in an investigation is prohibited.",
            "Concerns may be reported to a supervisor, Human Resources, or any manager the employee trusts.",
        ],
    )
    doc.add_paragraph(
        "Reports will be handled promptly, impartially, and as confidentially as possible while still allowing a fair investigation."
    )

    add_section_heading(doc, "4. Attendance, work hours, and pay", 1)
    add_bullets(
        doc,
        [
            "Employees should be on time, ready to work, and should notify a supervisor as soon as possible if they will be late or absent.",
            "All hours worked must be recorded accurately. Never ask another employee to record time for you.",
            "Nonexempt employees must receive pay for all hours worked and overtime must be paid in accordance with applicable law.",
            "Working unapproved overtime may lead to discipline, but it will still be paid when required by law.",
            "Meal and rest breaks, if offered, must be taken according to company procedure and local law.",
        ],
    )

    add_section_heading(doc, "5. Time off and leave", 1)
    doc.add_paragraph(
        "Full-time employees accrue 20 days of annual leave per year, beginning on the first day of employment. Annual leave is intended for rest, vacation, and personal time away from work."
    )
    add_bullets(
        doc,
        [
            "Annual leave accrues monthly at the rate of 1.67 days per month for full-time employees. Part-time and temporary employees do not accrue annual leave unless a written offer letter or local law says otherwise.",
            "Unused annual leave may be carried over up to 5 days into the next calendar year. Leave above the carryover limit is forfeited at year-end unless local law requires a different rule.",
            "Paid time off, sick leave, and holidays are provided according to the company's separate leave policy and local law.",
            "Eligible employees may have rights under the Family and Medical Leave Act (FMLA) or similar state laws for qualifying reasons.",
            "Employees should give notice as early as practical when they need leave and provide the information needed to determine whether the leave is protected.",
            "Medical information will be treated confidentially and shared only with people who need it to administer leave or comply with the law.",
        ],
    )

    add_section_heading(doc, "6. Workplace conduct", 1)
    add_bullets(
        doc,
        [
            "Treat coworkers, customers, vendors, and visitors respectfully.",
            "Use professional language and judgment in meetings, email, chat, and social media when speaking about the company or work matters.",
            "Follow instructions from supervisors, unless the instruction is unsafe or unlawful.",
            "Do not bring weapons, illegal drugs, or alcohol onto company property or work sites, except where local law and company policy clearly allow it.",
            "Report conflicts of interest, gifts, side work that conflicts with company interests, and other ethics concerns promptly.",
        ],
    )

    add_section_heading(doc, "7. Safety and health", 1)
    doc.add_paragraph(
        "The company will try to provide a safe workplace and expects employees to follow safety rules, training, and posted procedures."
    )
    add_bullets(
        doc,
        [
            "Report hazards, injuries, spills, or unsafe behavior immediately.",
            "Use required protective equipment and follow equipment instructions.",
            "Stop work and notify a supervisor if a task appears unsafe and you do not have the authority or training to correct it.",
            "Cooperate with injury reporting, safety investigations, and emergency drills.",
        ],
    )

    add_section_heading(doc, "8. Confidentiality, technology, and equipment", 1)
    add_bullets(
        doc,
        [
            "Company information, customer data, and employee records must be protected and shared only with authorized people.",
            "Use company devices, systems, and accounts responsibly and for business purposes unless local policy allows limited personal use.",
            "Do not install unauthorized software, disable security tools, or bypass approved controls.",
            "Return company property promptly when asked or when employment ends.",
        ],
    )

    add_section_heading(doc, "9. Remote and hybrid work", 1)
    add_bullets(
        doc,
        [
            "Remote work is a privilege based on business needs, not an automatic right.",
            "Remote employees must remain reachable during scheduled work hours, protect confidential information, and maintain a safe workspace.",
            "Timekeeping, performance, and conduct expectations are the same whether work is performed on-site or remotely.",
        ],
    )

    add_section_heading(doc, "10. Performance management and discipline", 1)
    doc.add_paragraph(
        "The company uses coaching, feedback, and corrective action to address performance or conduct issues. Depending on the situation, steps may include coaching, a written warning, a performance improvement plan, suspension, or termination."
    )
    add_bullets(
        doc,
        [
            "Serious misconduct may lead directly to termination.",
            "The company may skip or repeat steps depending on the facts, the law, and business needs.",
            "Nothing in this section limits the company's right to act immediately when necessary to protect people, property, or operations.",
        ],
    )

    add_section_heading(doc, "11. Separation of employment", 1)
    add_bullets(
        doc,
        [
            "When employment ends, employees should return company property, remove personal items, and complete any exit tasks.",
            "Final pay, accrued leave, and benefits will be handled according to law and company policy.",
            "Employees should keep private information confidential even after employment ends.",
        ],
    )

    add_section_heading(doc, "12. Employee acknowledgment", 1)
    doc.add_paragraph(
        "I acknowledge that I received or had access to the employee handbook, understand that it is a general policy guide and not a contract, and agree to follow the rules and procedures described in it."
    )
    ack_table = doc.add_table(rows=3, cols=2)
    ack_table.style = "Table Grid"
    set_table_widths(ack_table, [2.1, 4.5])
    labels = ["Employee name", "Signature", "Date"]
    for row, label in zip(ack_table.rows, labels):
        row.cells[0].text = label
        row.cells[1].text = ""
        set_cell_shading(row.cells[0], "F2F4F7")
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)

    add_section_heading(doc, "Appendix - Sources used to draft this sample", 1)
    doc.add_paragraph(
        "The following public sources informed the sample language. They are not reproduced verbatim here; they were used to shape the policies and reminders in this handbook."
    )
    add_source_entry(
        doc,
        "EEOC - Harassment",
        "https://www.eeoc.gov/harassment",
        "Defines unlawful harassment and emphasizes prevention, complaint procedures, and prompt corrective action.",
    )
    add_source_entry(
        doc,
        "EEOC - Employers",
        "https://www.eeoc.gov/employers",
        "Summarizes protected classes, harassment, reasonable accommodation, and retaliation protections.",
    )
    add_source_entry(
        doc,
        "DOL - FLSA employment relationship",
        "https://www.dol.gov/agencies/whd/fact-sheets/13-flsa-employment-relationship",
        "Summarizes minimum wage, overtime, recordkeeping, and child labor basics.",
    )
    add_source_entry(
        doc,
        "DOL - FMLA employee protections",
        "https://www.dol.gov/agencies/whd/fact-sheets/28a-fmla-employee-protections",
        "Explains job-protected leave, health benefits, and restoration rights for eligible employees.",
    )
    add_source_entry(
        doc,
        "OSHA - Employer responsibilities",
        "https://www.osha.gov/workers/employer-responsibilities",
        "Covers safe workplaces, hazard communication, and employee safety training.",
    )

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
